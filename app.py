from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import os
from io import BytesIO
import re

app = Flask(__name__)

ASSETS = "assets"
TEMPLATE_PATH = "label_template.docx"

SUBJECT_ICONS = {
    "Art and Design": "Art and Design Icon.png",
    "Brass Band": "Brass Band Icon.png",
    "Design and Technology": "Design and Technology Icon.png",
    "English": "English Icon.png",
    "Geography": "Geography Icon.png",
    "Guided Reading": "Guided Reading Icon.png",
    "History": "History Icon.png",
    "Homework": "Homework Icon.png",
    "Math": "Math Icon.png",
    "Music": "Music Icon.png",
    "Religious Education": "Religious Education Icon.png",
    "Science": "Science Icon.png",
    "Spanish": "Spanish Icon.png",
}


def format_name(name):
    name = str(name).strip()

    if "," in name:
        surname, firstname = name.split(",", 1)
        return f"{firstname.strip()} {surname.strip()}"

    return name


def chunk_list(data, size):
    for i in range(0, len(data), size):
        yield data[i:i + size]


def clean_students(raw_students):
    students = []

    for value in raw_students:
        name = str(value).strip()

        if not name or name.lower() == "nan":
            continue

        if name.lower().startswith("row count"):
            continue

        if name.lower() in ["name", "student", "student name", "students"]:
            continue

        students.append(name)

    return students


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)

    for table in list(cell.tables):
        table._element.getparent().remove(table._element)


def set_font(run, size, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "NTPreCursivefk"

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts

    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), "NTPreCursivefk")
    rFonts.set(qn("w:hAnsi"), "NTPreCursivefk")
    rFonts.set(qn("w:eastAsia"), "NTPreCursivefk")
    rFonts.set(qn("w:cs"), "NTPreCursivefk")


def name_font_size(name):
    length = len(name)

    if length <= 18:
        return 18
    elif length <= 24:
        return 16
    elif length <= 30:
        return 14
    else:
        return 12


def remove_problem_xml(element):
    # Duplicate bookmarks are a common cause of Word "unreadable content" repairs.
    for node in list(element.xpath(".//w:bookmarkStart")):
        node.getparent().remove(node)

    for node in list(element.xpath(".//w:bookmarkEnd")):
        node.getparent().remove(node)


def clear_document_body_but_keep_section(doc):
    body = doc._body._element
    sectPr = body.sectPr

    for child in list(body):
        if child is not sectPr:
            body.remove(child)


def append_before_sectpr(doc, element):
    body = doc._body._element
    sectPr = body.sectPr

    if sectPr is not None:
        body.insert(body.index(sectPr), element)
    else:
        body.append(element)


def add_page_break_before_next_table(doc):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    append_before_sectpr(doc, p)


def fill_label(cell, student, year_group, subject):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    logo_path = os.path.join(ASSETS, "STMP Logo.png")
    icon_path = os.path.join(ASSETS, SUBJECT_ICONS.get(subject, ""))

    p_logo = cell.add_paragraph()
    p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)

    run_logo = p_logo.add_run()
    if os.path.exists(logo_path):
        run_logo.add_picture(logo_path, width=Mm(18))

    p_name = cell.add_paragraph()
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)

    r_name = p_name.add_run(student)
    set_font(r_name, name_font_size(student), bold=True)

    p_subject = cell.add_paragraph()
    p_subject.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_subject.paragraph_format.space_before = Pt(0)
    p_subject.paragraph_format.space_after = Pt(0)

    r_subject = p_subject.add_run(subject)
    set_font(r_subject, 16)

    p_year = cell.add_paragraph()
    p_year.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_year.paragraph_format.space_before = Pt(0)
    p_year.paragraph_format.space_after = Pt(0)

    r_year = p_year.add_run(f"Year {year_group}")
    set_font(r_year, 16)

    p_icon = cell.add_paragraph()
    p_icon.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_icon.paragraph_format.space_before = Pt(0)
    p_icon.paragraph_format.space_after = Pt(0)

    run_icon = p_icon.add_run()
    if os.path.exists(icon_path):
        run_icon.add_picture(icon_path, width=Mm(18))


def fill_table(table, page_students, year_group, subject):
    label_positions = [
        (0, 0), (0, 2),
        (1, 0), (1, 2),
        (2, 0), (2, 2),
        (3, 0), (3, 2),
    ]

    for i, (r, c) in enumerate(label_positions):
        cell = table.cell(r, c)

        if i < len(page_students):
            student = format_name(page_students[i])
            fill_label(cell, student, year_group, subject)
        else:
            clear_cell(cell)


def build_docx(students, year_group, subject):
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError("label_template.docx was not found.")

    doc = Document(TEMPLATE_PATH)

    if not doc.tables:
        raise ValueError("The template must contain at least one label table.")

    section = doc.sections[0]
    section.top_margin = Mm(14)

    template_table_xml = deepcopy(doc.tables[0]._tbl)

    clear_document_body_but_keep_section(doc)

    pages = list(chunk_list(students, 8))

    for page_index, page_students in enumerate(pages):
        if page_index > 0:
            add_page_break_before_next_table(doc)

        table_xml = deepcopy(template_table_xml)
        remove_problem_xml(table_xml)

        append_before_sectpr(doc, table_xml)

        table = doc.tables[-1]
        fill_table(table, page_students, year_group, subject)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    subjects = sorted(SUBJECT_ICONS.keys())
    return render_template("index.html", subjects=subjects)


@app.route("/generate", methods=["POST"])
def generate():
    file = request.files["file"]
    year_group = request.form["year_group"]
    subject = request.form["subject"]

    df = pd.read_excel(file)
    raw_students = df.iloc[:, 0].dropna().tolist()
    students = clean_students(raw_students)

    docx_file = build_docx(students, year_group, subject)

    return send_file(
        docx_file,
        as_attachment=True,
        download_name=f"{subject}_Year{year_group}_Labels.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
